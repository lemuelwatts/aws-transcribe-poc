# Copyright 2025 Booz Allen Hamilton.
#
# Booz Allen Hamilton Confidential Information.
#
# The contents of this file are the intellectual property of
# Booz Allen Hamilton, Inc. ("BAH") and are subject to copyright protection
# under the laws of the United States and other countries.
#
# You acknowledge that misappropriation, misuse, or redistribution of content
# on the file could cause irreparable harm to BAH and/or to third parties.
#
# You may not copy, reproduce, distribute, publish, display, execute, modify,
# create derivative works of, transmit, sell or offer for resale, or in any way
# exploit any part of this code or program without BAH's express written permission.
#
# The contents of this code or program contains code
# that is itself or was created using artificial intelligence.
#
# To the best of our knowledge, this code does not infringe third-party intellectual
# property rights, contain errors, inaccuracies, bias, or security concerns.
#
# However, Booz Allen does not warrant, claim, or provide any implied
# or express warranty for the aforementioned, nor of merchantability
# or fitness for purpose.
#
# Booz Allen expressly limits liability, whether by contract, tort or in equity
# for any damage or harm caused by use of this artificial intelligence code or program.
#
# Booz Allen is providing this code or program "as is" with the understanding
# that any separately negotiated standards of performance for said code
# or program will be met for the duration of any applicable contract under which
# the code or program is provided.

"""Transcript processing router for normalizing AWS Transcribe output."""

import json
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel

from ..services.transcript_normalizer import NormalizedTranscript, TranscriptNormalizer

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path("output/normalized_transcripts")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

router = APIRouter(prefix="/transcript", tags=["transcript"])


def _format_timestamp(seconds: float) -> str:
    """Convert seconds to HH:MM:SS format."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def _generate_word_document(result: NormalizedTranscript, output_path: Path) -> None:
    """Generate a Word document with the transcript formatted by speaker.

    Args:
        result: The normalized transcript data.
        output_path: Path to save the Word document.
    """
    doc = Document()

    # Title
    title = doc.add_heading("Meeting Transcript", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Job: {result.job_name}")
    doc.add_paragraph(f"Speakers: {result.speakers_count}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()  # Blank line

    # Horizontal line
    doc.add_paragraph("─" * 50)

    # Transcript content
    for segment in result.segments:
        # Speaker header with timestamp
        timestamp_start = _format_timestamp(segment.start_time)
        timestamp_end = _format_timestamp(segment.end_time)

        speaker_para = doc.add_paragraph()
        speaker_run = speaker_para.add_run(
            f"{segment.speaker} [{timestamp_start} - {timestamp_end}]"
        )
        speaker_run.bold = True
        speaker_run.font.size = Pt(11)

        # Speaker's text
        text_para = doc.add_paragraph(segment.text)
        text_para.paragraph_format.left_indent = Inches(0.25)
        text_para.paragraph_format.space_after = Pt(12)

    doc.save(output_path)


class NormalizedTranscriptResponse(BaseModel):
    """Response model for normalized transcript."""

    success: bool
    job_name: str
    speakers_count: int
    segments_count: int
    saved_json_path: str
    saved_docx_path: str


@router.post("/normalize", response_model=NormalizedTranscriptResponse)
async def normalize_transcript(
    file: UploadFile = File(...),
) -> NormalizedTranscriptResponse:
    """Normalize an AWS Transcribe JSON file into speaker-grouped segments.

    Accepts an AWS Transcribe output JSON file and returns a normalized
    format where consecutive speech from the same speaker is merged into
    single segments.

    Args:
        file: The AWS Transcribe JSON output file.

    Returns:
        NormalizedTranscriptResponse with segments organized by speaker
        in chronological order.

    Raises:
        HTTPException: If the file is not valid JSON or missing required keys.
    """
    try:
        content = await file.read()
        data = json.loads(content.decode("utf-8"))
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid JSON file: {e!s}",
        )
    except UnicodeDecodeError as e:
        raise HTTPException(
            status_code=400,
            detail=f"File encoding error: {e!s}",
        )

    try:
        normalizer = TranscriptNormalizer()
        result = normalizer.normalize(data)

        # Generate output filenames with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{result.job_name}_normalized_{timestamp}"
        json_path = OUTPUT_DIR / f"{base_filename}.json"
        docx_path = OUTPUT_DIR / f"{base_filename}.docx"

        # Build JSON data
        json_data = {
            "job_name": result.job_name,
            "speakers_count": result.speakers_count,
            "segments": [
                {
                    "speaker": seg.speaker,
                    "start_time": seg.start_time,
                    "end_time": seg.end_time,
                    "text": seg.text,
                }
                for seg in result.segments
            ],
        }

        # Save JSON file
        with open(json_path, "w") as f:
            json.dump(json_data, f, indent=2)
        logger.info(f"Saved normalized JSON to {json_path}")

        # Generate and save Word document
        _generate_word_document(result, docx_path)
        logger.info(f"Saved Word document to {docx_path}")

        return NormalizedTranscriptResponse(
            success=True,
            job_name=result.job_name,
            speakers_count=result.speakers_count,
            segments_count=len(result.segments),
            saved_json_path=str(json_path),
            saved_docx_path=str(docx_path),
        )

    except ValueError as e:
        raise HTTPException(
            status_code=400,
            detail=str(e),
        )
